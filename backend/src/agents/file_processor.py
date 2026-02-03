"""
File Processing Tool using Vertex AI Gemini

Uses ADK agents that call Vertex AI Gemini to process PDFs/Excel files.
The file is wired into Gemini as a Part, and ADK handles the agent logic.

Supported file types:
- PDF (.pdf): Processed directly by Gemini's multimodal capabilities
- Excel (.xlsx, .xls): Read with openpyxl, converted to text for Gemini
- Word (.docx): Read with python-docx, converted to text for Gemini
"""

import base64
import io
from pathlib import Path
from typing import Any, Optional
import json

from google import genai
from google.genai import types


def get_gemini_client() -> genai.Client:
    """Create Vertex AI Gemini client."""
    # Uses GOOGLE_APPLICATION_CREDENTIALS or default credentials
    return genai.Client(vertexai=True)


def process_file_with_gemini(
    file_content: bytes,
    file_name: str,
    extraction_prompt: str,
    model: str = "gemini-2.0-flash"
) -> dict[str, Any]:
    """
    Process a file using Vertex AI Gemini.
    
    Args:
        file_content: Raw bytes of the uploaded file
        file_name: Original filename (used to determine MIME type)
        extraction_prompt: Instructions for what data to extract
        model: Gemini model to use
        
    Returns:
        Extracted data as a dictionary
    """
    client = get_gemini_client()
    
    # Determine MIME type from file extension
    mime_type = _get_mime_type(file_name)
    
    # For Excel files, convert to text representation first
    if mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     "application/vnd.ms-excel"]:
        text_content = _excel_to_text(file_content)
        parts = [
            types.Part.from_text(f"Excel file content:\n\n{text_content}"),
            types.Part.from_text(extraction_prompt)
        ]
    # For Word documents, convert to text
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text_content = _docx_to_text(file_content)
        parts = [
            types.Part.from_text(f"Word document content:\n\n{text_content}"),
            types.Part.from_text(extraction_prompt)
        ]
    # For PDFs and other files, use Gemini's native multimodal processing
    else:
        parts = [
            types.Part.from_bytes(data=file_content, mime_type=mime_type),
            types.Part.from_text(extraction_prompt)
        ]
    
    # Generate response with structured output
    response = client.models.generate_content(
        model=model,
        contents=types.Content(
            role="user",
            parts=parts
        ),
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.1  # Low temperature for accurate extraction
        )
    )
    
    # Parse JSON response
    try:
        return json.loads(response.text)
    except json.JSONDecodeError:
        # Return raw text if JSON parsing fails
        return {"raw_text": response.text, "parse_error": True}


def _get_mime_type(file_name: str) -> str:
    """Get MIME type from file extension."""
    ext = Path(file_name).suffix.lower()
    mime_types = {
        ".pdf": "application/pdf",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".csv": "text/csv",
    }
    return mime_types.get(ext, "application/octet-stream")


def _excel_to_text(file_content: bytes) -> str:
    """Convert Excel file to text representation for Gemini."""
    try:
        import openpyxl
        
        workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)
            
            text_parts.append("")  # Empty line between sheets
        
        return "\n".join(text_parts)
    except Exception as e:
        return f"Error reading Excel file: {e}"


def _docx_to_text(file_content: bytes) -> str:
    """Convert Word document to text for Gemini."""
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract tables
        for table in doc.tables:
            text_parts.append("\n--- Table ---")
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        return f"Error reading Word document: {e}"


# =============================================================================
# ADK Tool Definitions
# =============================================================================

ROTA_EXTRACTION_PROMPT = """
Analyze this file and extract rota/schedule information. Return a JSON object with:
{
    "startDate": "YYYY-MM-DD or null if not found",
    "endDate": "YYYY-MM-DD or null if not found",
    "staffAssignments": [
        {
            "staffName": "name",
            "staffId": "ID if found",
            "date": "YYYY-MM-DD",
            "shiftCode": "shift code like E, L, N, AL, etc."
        }
    ],
    "specialRequests": [
        {
            "staffName": "name",
            "date": "YYYY-MM-DD",
            "requestType": "leave/training/etc",
            "notes": "any notes"
        }
    ],
    "summary": "Brief description of what was found"
}

Extract as much information as possible. If certain fields are not present, use null.
"""

UNIT_EXTRACTION_PROMPT = """
Analyze this file and extract unit/staff configuration information. Return a JSON object with:
{
    "unitInfo": {
        "name": "unit name if found",
        "department": "department if found",
        "manager": "manager name if found"
    },
    "staff": [
        {
            "name": "full name",
            "staffId": "staff ID/code",
            "position": "Level 1/Level 2/Level 3 or similar",
            "type": "Direct Care or Non-Direct Care",
            "contractedHours": number or null,
            "comments": "any notes"
        }
    ],
    "shiftCodes": [
        {
            "code": "E, L, N, etc",
            "definition": "Early, Late, Night, etc",
            "description": "shift times",
            "hours": number,
            "type": "Direct Care or Non-Direct Care"
        }
    ],
    "summary": "Brief description of what was found"
}

Extract as much information as possible. If certain fields are not present, use null.
"""


def extract_rota_data(file_content: bytes, file_name: str) -> dict[str, Any]:
    """
    Extract rota/schedule data from an uploaded file.
    
    This is the main tool function called by the Rota Filling Agent.
    """
    return process_file_with_gemini(
        file_content=file_content,
        file_name=file_name,
        extraction_prompt=ROTA_EXTRACTION_PROMPT
    )


def extract_unit_data(file_content: bytes, file_name: str) -> dict[str, Any]:
    """
    Extract unit/staff configuration data from an uploaded file.
    
    This is the main tool function called by the Unit Filling Agent.
    """
    return process_file_with_gemini(
        file_content=file_content,
        file_name=file_name,
        extraction_prompt=UNIT_EXTRACTION_PROMPT
    )
